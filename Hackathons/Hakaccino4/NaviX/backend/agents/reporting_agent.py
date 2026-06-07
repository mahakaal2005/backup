import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from agents.base_agent import call_llm, make_thought
from config import REPORTS_DIR
from database import get_db
from datetime import datetime
import uuid

AGENT_NAME = "REPORTING_AGENT"

SYSTEM_PROMPT = """You are the Reporting Agent for Navix, a procurement risk analysis system.
Your role is to synthesize all risk analysis findings into a comprehensive, executive-ready report.

You will receive outputs from one or more specialist agents:
- Scheduler Agent (schedule variances and delivery delays)
- Political Risk Agent (geopolitical risks by country)
- Tariff Risk Agent (trade duty and customs risks)
- Logistics Risk Agent (shipping and transport disruptions)

Your job is to:
1. Create a well-structured executive summary
2. Combine all available risk findings cohesively
3. Provide a unified risk dashboard (overall risk level 1-5)
4. List prioritized recommendations (most critical first)
5. Identify interconnections between schedule delays and external risk factors

IMPORTANT FORMATTING RULES:
- Use Markdown tables wherever data lends itself to tabular display (item lists, country comparisons, risk summaries).
- For schedule data, always include a table like:
  | Item Name | Vendor Country | Expected Date | Actual Date | Delay (days) | Risk |
  |-----------|---------------|--------------|-------------|-------------|------|
- For country risk, include a summary table:
  | Country | Items | Avg Risk | Total Delay | Assessment |
  |---------|-------|----------|-------------|------------|
- Use bold text for risk levels and key figures.
- Use bullet points only for recommendations and narrative points, NOT for data that fits in a table.
- Do NOT include any image references, placeholder text like [Insert ...], or image markdown syntax.
- The interactive map is handled separately by the system — do not reference it.

Format your response as a professional Markdown report with these sections:
## Executive Summary
## Overall Risk Assessment
## Schedule Risk Analysis
## Country Risk Summary (table)
## Political Risk Analysis (if applicable)
## Tariff Risk Analysis (if applicable)
## Logistics Risk Analysis (if applicable)
## Key Findings & Interconnections
## Prioritized Recommendations
## Conclusion

Be concise, professional, and actionable. Skip sections where no data was provided."""


def _hex_to_rgb(hex_color: str):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color using OOXML shading."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10, color_hex: str = None, center: bool = False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color_hex:
        r, g, b = _hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_risk_dashboard(doc, metrics: dict):
    """Add a 3-column color-coded risk summary dashboard."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc.add_heading('Risk Dashboard', level=2)

    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    configs = [
        ('HIGH RISK', metrics.get('high_count', 0), metrics.get('high_pct', 0), 'C0392B', 'FADBD8'),
        ('MEDIUM RISK', metrics.get('medium_count', 0), metrics.get('medium_pct', 0), 'D35400', 'FAD7A0'),
        ('LOW RISK', metrics.get('low_count', 0), metrics.get('low_pct', 0), '1E8449', 'A9DFBF'),
    ]

    for col_idx, (label, count, pct, dark_hex, light_hex) in enumerate(configs):
        header_cell = table.cell(0, col_idx)
        _set_cell_bg(header_cell, dark_hex)
        _set_cell_text(header_cell, label, bold=True, font_size=10, color_hex='FFFFFF', center=True)

        value_cell = table.cell(1, col_idx)
        _set_cell_bg(value_cell, light_hex)
        value_cell.text = ''
        para = value_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_count = para.add_run(f'{count} items\n')
        run_count.bold = True
        run_count.font.size = Pt(14)
        run_pct = para.add_run(f'({pct}%)')
        run_pct.font.size = Pt(10)

    doc.add_paragraph()

    total_delay = metrics.get('total_delay_days', 0)
    total_items = metrics.get('high_count', 0) + metrics.get('medium_count', 0) + metrics.get('low_count', 0)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'Total Items Analyzed: ').bold = True
    summary_para.add_run(f'{total_items}   ')
    summary_para.add_run(f'Total Accumulated Delay: ').bold = True
    summary_para.add_run(f'{total_delay} days')
    doc.add_paragraph()


def _add_equipment_table(doc, items: list):
    """Add color-coded equipment schedule table."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading('Equipment Schedule Details', level=2)

    RISK_COLORS = {
        'High':   ('FADBD8', 'C0392B'),
        'Medium': ('FAD7A0', 'D35400'),
        'Low':    ('A9DFBF', '1E8449'),
    }

    headers = ['Item Name', 'Vendor Country', 'Expected Date', 'Actual Date', 'Variance', 'Status', 'Risk']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_row.cells[i], '2C3E50')
        _set_cell_text(hdr_row.cells[i], h, bold=True, font_size=9, color_hex='FFFFFF', center=True)

    for item in items:
        risk_cat = item.get('risk_category', 'Low')
        bg_color, text_color = RISK_COLORS.get(risk_cat, ('FFFFFF', '000000'))
        variance = item.get('variance_days', 0)
        variance_str = f'+{variance}d' if variance > 0 else f'{variance}d'

        row = table.add_row()
        cells = row.cells

        values = [
            item.get('item_name', ''),
            item.get('vendor_country', ''),
            item.get('expected_date', ''),
            item.get('actual_date', ''),
            variance_str,
            item.get('status', ''),
            risk_cat,
        ]

        for i, val in enumerate(values):
            _set_cell_bg(cells[i], bg_color)
            is_risk_col = (i == len(values) - 1)
            _set_cell_text(cells[i], val, bold=is_risk_col, font_size=9,
                          color_hex=text_color if is_risk_col else '2C3E50', center=is_risk_col)

    doc.add_paragraph()


def _add_country_table(doc, items: list):
    """Add a country risk summary table computed from item-level data."""
    from docx.shared import Pt

    countries: dict = {}
    for item in items:
        c = item.get('vendor_country', 'Unknown')
        if c not in countries:
            countries[c] = {'total': 0, 'delayed': 0, 'high': 0, 'total_variance': 0, 'risk_sum': 0}
        countries[c]['total'] += 1
        countries[c]['risk_sum'] += item.get('risk_level', 1)
        countries[c]['total_variance'] += max(item.get('variance_days', 0), 0)
        if item.get('status') in ('delayed', 'at_risk'):
            countries[c]['delayed'] += 1
        if item.get('risk_category') == 'High':
            countries[c]['high'] += 1

    doc.add_heading('Country Risk Breakdown', level=2)

    headers = ['Country', 'Items', 'Delayed', 'High Risk', 'Total Delay (days)', 'Avg Risk Score']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_row.cells[i], '2C3E50')
        _set_cell_text(hdr_row.cells[i], h, bold=True, font_size=9, color_hex='FFFFFF', center=True)

    sorted_countries = sorted(countries.items(), key=lambda x: -x[1]['risk_sum'])

    for country, stats in sorted_countries:
        avg = stats['risk_sum'] / stats['total'] if stats['total'] else 0
        risk_ratio = stats['high'] / stats['total'] if stats['total'] else 0

        if risk_ratio > 0.5 or stats['high'] >= 2:
            row_color = 'FADBD8'
        elif stats['delayed'] > 0 or stats['high'] > 0:
            row_color = 'FAD7A0'
        else:
            row_color = 'A9DFBF'

        row = table.add_row()
        values = [
            country,
            str(stats['total']),
            str(stats['delayed']),
            str(stats['high']),
            f"{stats['total_variance']}d",
            f"{avg:.1f}/5",
        ]
        for i, val in enumerate(values):
            _set_cell_bg(row.cells[i], row_color)
            _set_cell_text(row.cells[i], val, font_size=9, center=(i > 0))

    doc.add_paragraph()


def _parse_markdown_into_doc(doc, content: str):
    """Parse LLM markdown narrative into Word document paragraphs."""
    from docx.shared import Pt

    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            _add_inline_bold(para, text)
        elif stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
        else:
            para = doc.add_paragraph()
            _add_inline_bold(para, stripped)


def _add_inline_bold(para, text: str):
    """Split on **bold** markers and add runs accordingly."""
    from docx.shared import Pt
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def _build_docx_report(content: str, session_id: str, scheduler_metrics: dict = None) -> tuple:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        title_para = doc.add_heading('Navix — Procurement Risk Report', 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}  |  Session: {session_id[:8]}…")
        meta_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()

        has_metrics = (
            scheduler_metrics
            and scheduler_metrics.get('metrics')
            and scheduler_metrics.get('categorized_items')
        )

        if has_metrics:
            _add_risk_dashboard(doc, scheduler_metrics['metrics'])
            _add_equipment_table(doc, scheduler_metrics['categorized_items'])
            _add_country_table(doc, scheduler_metrics['categorized_items'])

            doc.add_heading('Narrative Analysis', level=1)

        _parse_markdown_into_doc(doc, content)

        report_id = str(uuid.uuid4())
        filename = f"navix_report_{report_id[:8]}.docx"
        file_path = os.path.join(REPORTS_DIR, filename)
        doc.save(file_path)

        blob_url = f"/reports/{filename}"
        return report_id, file_path, blob_url, filename

    except Exception as e:
        print(f"Error building docx: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None, None


def _save_report_to_db(report_id: str, session_id: str, file_path: str, blob_url: str):
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO reports (report_id, session_id, file_path, blob_url) VALUES (?, ?, ?, ?)",
            (report_id, session_id, file_path, blob_url)
        )
        conn.commit()
    finally:
        conn.close()


def run(
    message: str,
    scheduler_output: str = "",
    political_output: str = "",
    tariff_output: str = "",
    logistics_output: str = "",
    session_id: str = "",
    scheduler_metadata: dict = None,
) -> dict:
    thoughts = []

    inputs_available = []
    if scheduler_output:
        inputs_available.append("Scheduler")
    if political_output:
        inputs_available.append("Political Risk")
    if tariff_output:
        inputs_available.append("Tariff Risk")
    if logistics_output:
        inputs_available.append("Logistics Risk")

    thoughts.append(make_thought(
        AGENT_NAME, "synthesis_start",
        f"Beginning synthesis of {len(inputs_available)} agent output(s) into a comprehensive risk report.",
        output=f"Inputs: {', '.join(inputs_available) if inputs_available else 'None'}"
    ))

    context_parts = [f"User Query: {message}\n"]
    if scheduler_output:
        context_parts.append(f"## Schedule Risk Analysis\n{scheduler_output}\n")
    if political_output:
        context_parts.append(f"## Political Risk Analysis\n{political_output}\n")
    if tariff_output:
        context_parts.append(f"## Tariff Risk Analysis\n{tariff_output}\n")
    if logistics_output:
        context_parts.append(f"## Logistics Risk Analysis\n{logistics_output}\n")

    context = "\n".join(context_parts)

    thoughts.append(make_thought(
        AGENT_NAME, "report_generation",
        "Generating executive risk report combining all available agent findings.",
    ))

    response = call_llm(SYSTEM_PROMPT, context, max_tokens=1800)

    thoughts.append(make_thought(
        AGENT_NAME, "report_complete",
        "Risk report generated successfully.",
        output=response[:400] + "..." if len(response) > 400 else response
    ))

    report_content = None
    if session_id and response and "GROQ_API_KEY" not in response:

        sched_metrics = None
        if scheduler_metadata:
            metrics = scheduler_metadata.get("metrics", {})
            categorized = scheduler_metadata.get("metrics_full", {}).get("categorized", [])
            if metrics and categorized:
                sched_metrics = {
                    "metrics": metrics,
                    "categorized_items": categorized,
                }

        thoughts.append(make_thought(
            AGENT_NAME, "document_creation",
            "Saving report as a Word document (.docx) with visual tables and risk dashboard.",
        ))

        report_id, file_path, blob_url, filename = _build_docx_report(
            response, session_id, scheduler_metrics=sched_metrics
        )
        if report_id and blob_url:
            _save_report_to_db(report_id, session_id, file_path, blob_url)
            report_content = {
                "report_id": report_id,
                "blob_url": blob_url,
                "filename": filename,
            }
            thoughts.append(make_thought(
                AGENT_NAME, "document_saved",
                "Report saved successfully.",
                output=f"Report available at: {blob_url}"
            ))

    return {
        "response": response,
        "thoughts": thoughts,
        "report_content": report_content,
    }
